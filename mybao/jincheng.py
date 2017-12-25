import os
import random
import time
from multiprocessing import Pool


def task(name):
    print("run task %s,pid(%s)" % (name, os.getpid()))
    start = time.time()
    time.sleep(random.random() * 3)
    print("task %s runs %0.2f seconds" % (name, time.time() - start))


from docx import Document
from contextlib import contextmanager


@contextmanager
def scw(s, j, o):
    if not os.path.exists("d:\\ubuntu/b%s.docx" % o):
        document = Document()
    else:
        document = Document(docx="d:\\ubuntu/b%s.docx" % o)

    yield document
    document.save("d:\\ubuntu/b%s.docx" % o)
    print("word写入完成")


if __name__ == '__main__':
    print("fu jincheng run %s" % os.getpid())
    p = Pool(4)
    for i in "临兵斗者皆阵列在前":
        p.apply_async(task, args=(i,))
    p.close()
    p.join()
    print("all done")
    with scw("u", "i", "o") as d:
        d.add_heading("天苍苍野茫茫", level=2)
